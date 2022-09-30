import os
import subprocess
from weakref import WeakValueDictionary
import webbrowser
from mailmerge import MailMerge
from datetime import date
from docx import Document
from docx.shared import Inches
import fitz
import qrcode


def fill_the_doc(number, serial_number, modification, control_test_0, control_test_1):
    document = MailMerge('protocol_after_psi/protocol_template.docx')

    document.merge(
        Protocol_Number=number,
        Protocol_Date='{:%d.%m.%Y}'.format(date.today()),
        Serial_Number=serial_number,
        Date='{:%d.%m.%Y}'.format(date.today()),
        Control_Test_Data_0=control_test_0,
        Control_Test_Data_1=control_test_1,
        Modification=modification,
    )
    date_today = '{:%d-%m-%Y}'.format(date.today())
    input_docx = f'{serial_number}-{date_today}'
    document.write(input_docx)

    return input_docx


def convert_docx_to_pdf(input_docx):
    # libre_office = r'C:/Program Files/LibreOffice/program/soffice.exe'
    libre_office = r'/usr/bin/libreoffice'
    process1 = subprocess.Popen([libre_office, '--headless', '--convert-to', 'pdf', '--outdir', 'protocol_after_psi', input_docx])
    process1.communicate()
    os.remove(f'{input_docx}')


def print_sticker(serial_number, modification):
    img = qrcode.make(serial_number)
    qr_name = f'stickers/{serial_number}.png'
    img.save(qr_name)

    document = MailMerge('platan/programs/stickers/sticker_template.docx')

    document.merge(
        Protocol_Date='{:%d.%m.%Y}'.format(date.today()),
        Serial_Num=serial_number,
        Date='{:%d.%m.%Y}'.format(date.today()),
        Modification=modification,
    )
    date_today = '{:%d-%m-%Y}'.format(date.today())
    input_docx = f'{serial_number}-{date_today}'
    document.write(input_docx)
    doc = Document(f'{input_docx}')
    tables = doc.tables
    p = tables[0].rows[8].cells[1].add_paragraph()
    r = p.add_run()
    r.add_picture(f'{qr_name}', width=Inches(.9), height=Inches(.9))
    doc.save(f'{input_docx}')
    os.remove(f'{qr_name}')
    convert_docx_to_pdf_package(input_docx)


def convert_docx_to_pdf_package(input_docx):
    # libre_office = r'C:/Program Files/LibreOffice/program/soffice.exe'
    libre_office = r'/usr/bin/libreoffice'
    process2 = subprocess.Popen([libre_office, '--headless', '--convert-to', 'pdf', '--outdir', 'stickers', input_docx])
    process2.communicate()
    os.remove(f'{input_docx}')
    input_file = f'stickers/{input_docx}.pdf'
    output_file = f'stickers/{input_docx}_print.pdf'
    file_handle = fitz.open(input_file)
    pages_list = [0]
    file_handle.select(pages_list)
    file_handle.save(output_file)
    os.remove(input_file)
    
    # browser = webbrowser.get('Firefox')
    # browser.open_new_tab(f'{output_file}')
    # subprocess.run(['evince', f'{output_file}'], stdout=subprocess.PIPE).stdout.decode('utf-8')


def print_sticker_passport(serial_number):
    document = MailMerge('platan/programs/stickers/passport_template.docx')

    document.merge(
        Serial_Num=serial_number,
    )
    date_today = '{:%d-%m-%Y}'.format(date.today())
    input_docx_passport = f'{serial_number}-{date_today}-passport'
    document.write(input_docx_passport)
    convert_docx_to_pdf_package(input_docx_passport)


def convert_docx_to_pdf_passport(input_docx_passport):
    # libre_office = r'C:/Program Files/LibreOffice/program/soffice.exe'
    libre_office = r'/usr/bin/libreoffice'
    process3 = subprocess.Popen([libre_office, '--headless', '--convert-to', 'pdf', '--outdir', 'stickers', input_docx_passport])
    process3.communicate()
    os.remove(f'{input_docx_passport}')
    input_file = f'platan/programs/stickers/{input_docx_passport}.pdf'
    output_file = f'platan/programs/stickers/{input_docx_passport}_print.pdf'
    file_handle = fitz.open(input_file)
    pages_list = [0]
    file_handle.select(pages_list)
    file_handle.save(output_file)
    os.remove(input_file)
    
    # browser = webbrowser.get('Firefox')
    # browser.open_new_tab(f'{output_file}')
    # subprocess.run(['evince', f'{output_file}'], stdout=subprocess.PIPE).stdout.decode('utf-8')
